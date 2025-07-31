"""
HackRx 6.0 - Enhanced Document Processor
Supports PDF, DOCX, and Email document processing
"""

import aiohttp
import PyPDF2
import docx
import email
import email.mime.text
from email.mime.multipart import MIMEMultipart
import re
import io
import asyncio
from typing import Optional, Dict, Any, List, Tuple
import logging
from urllib.parse import urlparse
import mimetypes

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Enhanced document processor supporting multiple file formats
    """
    
    def __init__(self):
        self.supported_formats = {
            'pdf': ['.pdf'],
            'docx': ['.docx', '.doc'],
            'email': ['.eml', '.msg', '.email', '.mbox'],
            'text': ['.txt']
        }
        self.is_initialized = False
        self.processing_stats = {
            'pdf_processed': 0,
            'docx_processed': 0,
            'email_processed': 0,
            'total_processed': 0,
            'failed_processing': 0
        }
    
    async def initialize(self):
        """Initialize the document processor"""
        self.is_initialized = True
        print("✅ Enhanced Document Processor initialized (PDF + DOCX + Email)")
    
    async def process_document(self, document_url: str) -> Optional[str]:
        """
        Process document from URL - auto-detects format and uses appropriate processor
        """
        try:
            print(f"📄 Processing document: {document_url[:100]}...")
            
            # Download document
            document_data = await self._download_document(document_url)
            
            if not document_data:
                raise Exception("Failed to download document")
            
            # Detect document type
            doc_type = self._detect_document_type(document_url, document_data)
            print(f"🔍 Detected document type: {doc_type}")
            
            # Process based on type
            if doc_type == 'pdf':
                text = await self._process_pdf(document_data)
                self.processing_stats['pdf_processed'] += 1
            elif doc_type == 'docx':
                text = await self._process_docx(document_data)
                self.processing_stats['docx_processed'] += 1
            elif doc_type == 'email':
                text = await self._process_email(document_data)
                self.processing_stats['email_processed'] += 1
            elif doc_type == 'text':
                text = document_data.decode('utf-8', errors='ignore')
            else:
                # Fallback: try as text first, then PDF
                try:
                    text = document_data.decode('utf-8', errors='ignore')
                    if len(text.strip()) < 100:  # Too short, probably not text
                        text = await self._process_pdf(document_data)
                except:
                    text = await self._process_pdf(document_data)
            
            if text and len(text.strip()) > 50:
                # Clean and optimize text
                cleaned_text = self._clean_text(text)
                self.processing_stats['total_processed'] += 1
                
                print(f"✅ Document processed successfully - {len(cleaned_text)} characters extracted")
                return cleaned_text
            else:
                raise Exception("Insufficient content extracted from document")
                
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            self.processing_stats['failed_processing'] += 1
            return None
    
    async def _download_document(self, url: str) -> Optional[bytes]:
        """Download document from URL"""
        try:
            timeout = aiohttp.ClientTimeout(total=30)
            
            async with aiohttp.ClientSession(timeout=timeout) as session:
                async with session.get(url) as response:
                    if response.status == 200:
                        content = await response.read()
                        print(f"📥 Downloaded {len(content)} bytes")
                        return content
                    else:
                        raise Exception(f"HTTP {response.status}: {response.reason}")
                        
        except Exception as e:
            logger.error(f"Document download failed: {e}")
            return None
    
    def _detect_document_type(self, url: str, data: bytes) -> str:
        """Detect document type from URL and content"""
        # Check URL extension first
        parsed_url = urlparse(url)
        path = parsed_url.path.lower()
        
        for doc_type, extensions in self.supported_formats.items():
            if any(path.endswith(ext) for ext in extensions):
                return doc_type
        
        # Check content headers/magic bytes
        if data.startswith(b'%PDF'):
            return 'pdf'
        elif data.startswith(b'PK\x03\x04') or data.startswith(b'PK\x05\x06'):
            # ZIP-based formats (DOCX)
            return 'docx'
        elif b'Message-ID:' in data[:1000] or b'From:' in data[:500]:
            return 'email'
        
        # Default fallback
        return 'text'
    
    async def _process_pdf(self, pdf_data: bytes) -> str:
        """Extract text from PDF data"""
        try:
            pdf_file = io.BytesIO(pdf_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Add page marker for better context
                        text_content.append(f"\n--- Page {page_num + 1} ---\n")
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {e}")
                    continue
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise Exception("No text content extracted from PDF")
            
            print(f"📄 PDF processed: {len(pdf_reader.pages)} pages, {len(full_text)} characters")
            return full_text
            
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise
    
    async def _process_docx(self, docx_data: bytes) -> str:
        """Extract text from DOCX data"""
        try:
            docx_file = io.BytesIO(docx_data)
            document = docx.Document(docx_file)
            
            text_content = []
            
            # Extract paragraph text
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract table text
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise Exception("No text content extracted from DOCX")
            
            print(f"📄 DOCX processed: {len(document.paragraphs)} paragraphs, {len(document.tables)} tables, {len(full_text)} characters")
            return full_text
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise
    
    async def _process_email(self, email_data: bytes) -> str:
        """Extract text from email data"""
        try:
            # Try to decode as email
            if isinstance(email_data, bytes):
                email_content = email_data.decode('utf-8', errors='ignore')
            else:
                email_content = email_data
            
            # Parse email
            msg = email.message_from_string(email_content)
            
            text_content = []
            
            # Extract headers
            text_content.append("--- Email Headers ---")
            for header in ['From', 'To', 'Subject', 'Date']:
                if msg.get(header):
                    text_content.append(f"{header}: {msg.get(header)}")
            
            text_content.append("\n--- Email Body ---")
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == 'text/plain':
                        payload = part.get_payload(decode=True)
                        if payload:
                            try:
                                text_content.append(payload.decode('utf-8', errors='ignore'))
                            except:
                                text_content.append(str(payload))
                    elif content_type == 'text/html':
                        # Simple HTML stripping
                        payload = part.get_payload(decode=True)
                        if payload:
                            try:
                                html_text = payload.decode('utf-8', errors='ignore')
                                clean_text = re.sub(r'<[^>]+>', ' ', html_text)
                                text_content.append(clean_text)
                            except:
                                pass
            else:
                # Single part email
                payload = msg.get_payload(decode=True)
                if payload:
                    try:
                        text_content.append(payload.decode('utf-8', errors='ignore'))
                    except:
                        text_content.append(str(payload))
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise Exception("No text content extracted from email")
            
            print(f"📧 Email processed: {len(full_text)} characters")
            return full_text
            
        except Exception as e:
            logger.error(f"Email processing failed: {e}")
            # Fallback: treat as plain text
            try:
                return email_data.decode('utf-8', errors='ignore')
            except:
                raise Exception("Failed to process email document")
    
    def _clean_text(self, text: str) -> str:
        """Clean and optimize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page breaks and form feeds
        text = re.sub(r'[\f\r]+', '\n', text)
        
        # Normalize line breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove very short lines (often artifacts)
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if len(line) > 2:  # Keep lines with more than 2 characters
                cleaned_lines.append(line)
        
        # Rejoin and clean
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove excessive repeated characters
        cleaned_text = re.sub(r'(.)\1{4,}', r'\1\1\1', cleaned_text)
        
        # Final cleanup
        cleaned_text = cleaned_text.strip()
        
        return cleaned_text
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get document processing statistics"""
        total = self.processing_stats['total_processed']
        failed = self.processing_stats['failed_processing']
        
        return {
            'initialized': self.is_initialized,
            'formats_supported': list(self.supported_formats.keys()),
            'pdf_processed': self.processing_stats['pdf_processed'],
            'docx_processed': self.processing_stats['docx_processed'],
            'email_processed': self.processing_stats['email_processed'],
            'total_processed': total,
            'failed_processing': failed,
            'success_rate': (total / (total + failed) * 100) if (total + failed) > 0 else 0.0
        }